from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = (
    ROOT
    / "quality"
    / "manual-tests"
    / "stage-d1"
    / "ARTestCLI_Manual_Test_Report_Native_Extensions_v1.0.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "B8C2CC"
MUTED = RGBColor(89, 97, 107)


def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text, *, bold=False, color="000000", align=None, size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    set_font(paragraph.add_run(text), size=size, bold=bold, color=color)


def style_table(table, header=True):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=9.5)
        if header and row_index == 0:
            set_repeat_header(row)
            for cell in row.cells:
                shade(cell, HEADER_FILL)
                for run in cell.paragraphs[0].runs:
                    set_font(run, size=9.5, bold=True, color=NAVY)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color="666666")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_metadata(document, rows):
    table = document.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=NAVY, size=10)
        set_cell_text(cells[1], value, size=10)
        shade(cells[0], HEADER_FILL)
    set_table_geometry(table, [1700, 7660])
    style_table(table, header=False)
    document.add_paragraph()


def add_note(document, label, text):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run(label + ": "), size=10, bold=True, color=NAVY)
    set_font(paragraph.add_run(text), size=10)
    document.add_paragraph()


def add_steps(document, rows):
    table = document.add_table(rows=1, cols=3)
    headers = ("Step", "Action", "Expected result")
    for index, text in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            text,
            bold=True,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for step, action, expected in rows:
        cells = table.add_row().cells
        set_cell_text(
            cells[0], str(step), bold=True, color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], action)
        set_cell_text(cells[2], expected)
    set_table_geometry(table, [600, 4380, 4380])
    style_table(table)
    document.add_paragraph()


def add_evidence_block(document):
    table = document.add_table(rows=4, cols=2)
    fields = (
        ("Actual result", ""),
        ("Evidence references", ""),
        ("Verdict", "Pending / Passed / Failed / Blocked"),
        ("Defect ID / Notes", ""),
    )
    for row, (label, value) in zip(table.rows, fields):
        set_cell_text(row.cells[0], label, bold=True, color=NAVY, size=9.5)
        set_cell_text(row.cells[1], value, size=9.5)
        shade(row.cells[0], LIGHT_FILL)
    set_table_geometry(table, [1800, 7560])
    style_table(table, header=False)


TEST_CASES = [
    {
        "id": "MT-D1-001",
        "title": "Official Debug and Release regression workflow",
        "objective": "Confirm the solution, ABI fingerprint, Google Test suite, and validated HTML reports complete in both supported configurations.",
        "preconditions": "Visual Studio 18 Insiders is installed at the documented path. Run from the repository root.",
        "steps": [
            (1, "Run .\\scripts\\build.ps1 -Configuration Debug -Platform x64.", "Build succeeds; ABI contract executable returns 0; Google Test reports 49 passed and 0 failed."),
            (2, "Open artifacts\\test-results\\x64\\Debug\\ARTestCLI.UnitTests.html.", "Overall, aggregate counters, and all test-case verdicts are PASSED and consistent."),
            (3, "Run .\\scripts\\build.ps1 -Configuration Release -Platform x64.", "Release produces the same 49/49 passed baseline and a validated HTML report."),
            (4, "Record both HTML reports and console output as evidence.", "Evidence identifies configuration, platform, test count, and final verdict."),
        ],
    },
    {
        "id": "MT-D1-002",
        "title": "Native command-to-driver execution through ARTestEngine.dll",
        "objective": "Verify the public facade and C ABI load both packages and execute a command through a driver service.",
        "preconditions": "Release build from MT-D1-001 completed successfully.",
        "steps": [
            (1, "Set $cli = '.\\artifacts\\bin\\x64\\Release\\ARTestCLI.exe' and $extensions = '.\\artifacts\\extensions\\x64\\Release'.", "Both paths exist."),
            (2, "Run & $cli extension-run '.\\source\\Scripts\\ExtensionScript.json' $extensions.", "Process returns exit code 0."),
            (3, "Inspect the Engine event lines.", "They show catalog load, INITIALIZING, driver initialization, RUNNING, command service completion, CLEANING_UP, driver shutdown, and COMPLETED in that order."),
            (4, "Inspect the final JSON line.", "status is passed; plannedSteps, executedSteps, and passedSteps are 1; the command ID is com.artest.command.sample.power-cycle."),
        ],
    },
    {
        "id": "MT-D1-003",
        "title": "Approved extension root validation",
        "objective": "Confirm the Engine rejects a missing or non-directory extension root without running a plan.",
        "preconditions": "Release build completed.",
        "steps": [
            (1, "Run & $cli extension-run '.\\source\\Scripts\\ExtensionScript.json' '.\\path-that-does-not-exist'.", "Process returns exit code 5."),
            (2, "Inspect stderr.", "It reports EXTENSION_ROOT_INVALID and identifies the requested root."),
            (3, "Confirm no INITIALIZING or RUNNING event was emitted.", "No instrument or command component was created."),
        ],
    },
    {
        "id": "MT-D1-004",
        "title": "Incompatible ABI rejection before DLL loading",
        "objective": "Verify manifest compatibility is evaluated before the native entry is required or loaded.",
        "preconditions": "Use the checked-in incompatible fixture; its entry DLL intentionally does not exist.",
        "steps": [
            (1, "Run & $cli extension-run '.\\source\\Scripts\\ExtensionScript.json' '.\\quality\\manual-tests\\stage-d1\\data\\incompatible'.", "Process returns exit code 5."),
            (2, "Inspect stderr.", "It reports EXTENSION_RUNTIME_INCOMPATIBLE and the BadAbi manifest path."),
            (3, "Confirm the message is not EXTENSION_LOAD_FAILED.", "Compatibility rejection happened before LoadLibrary."),
        ],
    },
    {
        "id": "MT-D1-005",
        "title": "Corrupt manifest containment",
        "objective": "Confirm malformed JSON is contained at the Engine boundary with an actionable diagnostic.",
        "preconditions": "Use the checked-in corrupt manifest fixture.",
        "steps": [
            (1, "Run & $cli extension-run '.\\source\\Scripts\\ExtensionScript.json' '.\\quality\\manual-tests\\stage-d1\\data\\corrupt'.", "Process returns exit code 5."),
            (2, "Inspect stderr.", "It reports EXTENSION_CATALOG_EXCEPTION with JSON parse context and the approved root."),
            (3, "Confirm ARTestCLI remains responsive and can run MT-D1-002 afterward.", "A corrupt package does not damage the valid Release package directory."),
        ],
    },
    {
        "id": "MT-D1-006",
        "title": "Cooperative Ctrl+C cancellation with guaranteed cleanup",
        "objective": "Verify Ctrl+C reaches the Engine session and cleanup shuts down the active driver.",
        "preconditions": "Release build completed. Use an interactive PowerShell console.",
        "steps": [
            (1, "Run & $cli extension-run '.\\quality\\manual-tests\\stage-d1\\data\\cancel-extension.json' $extensions.", "The command enters its 30-second hold after RUNNING."),
            (2, "After the step starts, press Ctrl+C once.", "The process does not terminate abruptly; cooperative cancellation is requested."),
            (3, "Inspect the event sequence.", "It shows CANCELLING, CLEANING_UP, 'Simulated power driver shut down.', instrument shutdown, and CANCELLED."),
            (4, "Inspect the final JSON and $LASTEXITCODE.", "status is cancelled and the process exit code is 5."),
        ],
    },
    {
        "id": "MT-D1-007",
        "title": "Cleanup failure overrides a passing command",
        "objective": "Verify driver shutdown failure remains visible and changes the overall verdict.",
        "preconditions": "Release build completed.",
        "steps": [
            (1, "Run & $cli extension-run '.\\quality\\manual-tests\\stage-d1\\data\\cleanup-failure-extension.json' $extensions.", "The command step itself completes successfully."),
            (2, "Inspect cleanup diagnostics and state events.", "Shutdown is attempted; a simulated shutdown failure is reported; terminal state is FAILED."),
            (3, "Inspect the final JSON and $LASTEXITCODE.", "status is error, failureKind is 3 (Cleanup), and exit code is 5."),
        ],
    },
]


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("ARTest Quality | Stage D1 Manual Validation"), size=9, color="666666")
    add_page_number(section.footer.paragraphs[0])

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_font(kicker.add_run("QUALITY ASSURANCE"), size=10, bold=True, color=BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_font(title.add_run("ARTestCLI Manual Test Report"), size=24, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(
        subtitle.add_run("Stage D1 - Native Extension Vertical Slice"),
        size=14,
        color=DARK_BLUE,
    )

    add_metadata(
        document,
        (
            ("Document ID", "ARTEST-QA-D1-MANUAL-001"),
            ("Version", "1.0"),
            ("Status", "Pending manual execution"),
            ("Prepared for", "ARTestCLI / ARTestEngine Stage D1 acceptance"),
            ("Tester", ""),
            ("Execution date", ""),
            ("Commit / branch", ""),
        ),
    )

    document.add_heading("Purpose and acceptance scope", level=1)
    document.add_paragraph(
        "This report captures repeatable manual evidence for the trusted-native "
        "Stage D1 vertical slice. It validates the public Engine boundary, "
        "manifest-first discovery, Command Plugin to Instrument Driver routing, "
        "cooperative cancellation, cleanup, diagnostics, and generated reports."
    )
    add_note(
        document,
        "Acceptance rule",
        "All seven manual cases must be Passed, and the Debug and Release "
        "automated reports must each show 49 passed, 0 failed, and 0 skipped.",
    )

    document.add_page_break()
    document.add_heading("Execution environment", level=1)
    add_metadata(
        document,
        (
            ("Operating system", "Windows:"),
            ("Visual Studio", "18 Insiders / v145"),
            ("Platform", "x64"),
            ("Configurations", "Debug and Release"),
            ("Repository", "D:\\GitHub\\main\\ARTestCLI"),
            ("PowerShell version", ""),
            ("Machine / tester", ""),
        ),
    )

    document.add_heading("Test summary", level=1)
    summary = document.add_table(rows=1, cols=4)
    for index, text in enumerate(("ID", "Scenario", "Expected verdict", "Actual verdict")):
        set_cell_text(summary.rows[0].cells[index], text, bold=True, color=NAVY)
    for case in TEST_CASES:
        cells = summary.add_row().cells
        set_cell_text(cells[0], case["id"], bold=True, color=NAVY)
        set_cell_text(cells[1], case["title"])
        set_cell_text(cells[2], "Passed", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], "Pending", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(summary, [1150, 5010, 1600, 1600])
    style_table(summary)

    document.add_page_break()
    for index, case in enumerate(TEST_CASES):
        document.add_heading(f'{case["id"]} | {case["title"]}', level=1)
        add_metadata(
            document,
            (
                ("Objective", case["objective"]),
                ("Preconditions", case["preconditions"]),
                ("Expected verdict", "Passed"),
            ),
        )
        document.add_heading("Procedure and expected results", level=2)
        add_steps(document, case["steps"])
        document.add_heading("Execution evidence", level=2)
        add_evidence_block(document)
        if index != len(TEST_CASES) - 1:
            document.add_page_break()

    document.add_page_break()
    document.add_heading("Final acceptance", level=1)
    add_metadata(
        document,
        (
            ("Total manual cases", str(len(TEST_CASES))),
            ("Passed", ""),
            ("Failed", ""),
            ("Blocked / Not run", ""),
            ("Automated Debug verdict", ""),
            ("Automated Release verdict", ""),
            ("Overall Stage D1 verdict", "Pending"),
        ),
    )
    document.add_heading("Approval", level=2)
    approval = document.add_table(rows=3, cols=3)
    for index, text in enumerate(("Role", "Name / Signature", "Date")):
        set_cell_text(approval.rows[0].cells[index], text, bold=True, color=NAVY)
    for row_index, role in enumerate(("Tester", "Technical reviewer"), start=1):
        set_cell_text(approval.rows[row_index].cells[0], role, bold=True)
        set_cell_text(approval.rows[row_index].cells[1], "")
        set_cell_text(approval.rows[row_index].cells[2], "")
    set_table_geometry(approval, [1900, 5060, 2400])
    style_table(approval)

    core = document.core_properties
    core.title = "ARTestCLI Manual Test Report - Stage D1"
    core.subject = "Native extension vertical slice acceptance"
    core.author = "ARTest Quality"
    core.keywords = "ARTestCLI, ARTestEngine, Stage D1, manual testing"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
