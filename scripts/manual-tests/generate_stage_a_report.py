from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = (
    REPOSITORY_ROOT
    / "quality"
    / "manual-tests"
    / "stage-a"
    / "ARTestCLI_Manual_Test_Report_Safe_Base_v1.0.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5D6B79"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WHITE = "FFFFFF"


TEST_CASES = [
    {
        "id": "MT-A-001",
        "title": "Compilación y regresión Debug x64",
        "objective": "Confirmar que la solución compila con C++20, v145 y /W4, y que la línea base automatizada es estable.",
        "preconditions": "Visual Studio 18 Insiders instalado en D:\\Program Files\\Microsoft Visual Studio\\18\\Insiders.",
        "steps": [
            ("1", "Abrir PowerShell en la raíz D:\\GitHub\\main\\ARTestCLI."),
            ("2", "Ejecutar: .\\scripts\\build.ps1 -Configuration Debug -Platform x64"),
            ("3", "Esperar a que finalicen MSBuild, Google Test y el generador HTML."),
            ("4", "Registrar el valor de $LASTEXITCODE y adjuntar captura de la salida."),
        ],
        "expected": "Código 0; compilación sin warnings; 15/15 pruebas PASSED; se generan XML y HTML en artifacts\\test-results\\x64\\Debug.",
    },
    {
        "id": "MT-A-002",
        "title": "Regresión Release x64",
        "objective": "Validar la configuración utilizada para una entrega o integración.",
        "preconditions": "MT-A-001 aprobado.",
        "steps": [
            ("1", "Ejecutar: .\\scripts\\build.ps1 -Configuration Release -Platform x64"),
            ("2", "Confirmar que se ejecuta la validación sintética del generador HTML."),
            ("3", "Registrar $LASTEXITCODE y abrir el reporte HTML Release."),
        ],
        "expected": "Código 0; 15/15 pruebas PASSED; resumen Overall PASSED, Failed 0 y Skipped 0.",
    },
    {
        "id": "MT-A-003",
        "title": "Descubrimiento en Visual Studio Test Explorer",
        "objective": "Confirmar que las pruebas forman parte de la solución y son ejecutables por un developer.",
        "preconditions": "Abrir source\\ARTestCLI.sln con Visual Studio Insiders; x64 y Debug.",
        "steps": [
            ("1", "Seleccionar Build > Build Solution."),
            ("2", "Abrir Test > Test Explorer."),
            ("3", "Esperar al descubrimiento y seleccionar Run All Tests."),
            ("4", "Adjuntar captura con las suites y el veredicto."),
        ],
        "expected": "Test Explorer descubre 15 pruebas en 4 suites y todas terminan Passed.",
    },
    {
        "id": "MT-A-004",
        "title": "Compilación offline de un script válido",
        "objective": "Demostrar que compile valida el documento sin inicializar recursos de hardware.",
        "preconditions": "Build Debug aprobado.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli compile '.\\source\\Scripts\\TestScript.json'"),
            ("3", "Registrar $LASTEXITCODE y revisar toda la salida."),
        ],
        "expected": "Código 0 y mensaje 'Valid script. No instruments were initialized.'; no aparecen mensajes de inicialización, ejecución o apagado.",
    },
    {
        "id": "MT-A-005",
        "title": "Ejecución válida y ciclo de vida de instrumentos",
        "objective": "Comprobar inicialización explícita, ejecución ordenada, resultado global y apagado.",
        "preconditions": "Build Debug aprobado; los instrumentos actuales son simulados y no requieren hardware real.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\source\\Scripts\\TestScript.json'"),
            ("3", "Registrar $LASTEXITCODE y conservar la salida completa."),
        ],
        "expected": "Código 0; se inicializan PS1 y CAN1; se ejecutan los pasos 1 a 4; el resultado es PASSED; ambos instrumentos se apagan.",
    },
    {
        "id": "MT-A-006",
        "title": "JSON corrupto y versión no soportada",
        "objective": "Verificar que documentos inválidos son rechazados antes de construir o inicializar instrumentos.",
        "preconditions": "Fixtures disponibles en quality\\manual-tests\\stage-a\\data.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli compile '.\\quality\\manual-tests\\stage-a\\data\\malformed-json.json'; $LASTEXITCODE"),
            ("3", "Confirmar código 3 y un diagnóstico de JSON inválido."),
            ("4", "& $cli compile '.\\quality\\manual-tests\\stage-a\\data\\unsupported-version.json'; $LASTEXITCODE"),
            ("5", "Confirmar código 3 y un diagnóstico de versión no soportada."),
        ],
        "expected": "Ambas ejecuciones devuelven 3, explican la causa y no inicializan instrumentos.",
    },
    {
        "id": "MT-A-007",
        "title": "Rechazo atómico de comando desconocido",
        "objective": "Evitar que una secuencia parcial se ejecute cuando contiene un comando no registrado.",
        "preconditions": "Fixture unknown-command.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-a\\data\\unknown-command.json'"),
            ("3", "Registrar $LASTEXITCODE y revisar la salida."),
        ],
        "expected": "Código 3 y diagnóstico de comando desconocido; no se ejecuta ningún paso ni se inicializa hardware.",
    },
    {
        "id": "MT-A-008",
        "title": "Falla controlada de inicialización",
        "objective": "Confirmar que una falla de hardware/configuración tiene código propio y no continúa con la secuencia.",
        "preconditions": "Fixture initialization-failure.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-a\\data\\initialization-failure.json'"),
            ("3", "Registrar $LASTEXITCODE y revisar el diagnóstico."),
        ],
        "expected": "Código 4; diagnóstico de recurso hw-rsrc faltante; no se ejecuta el paso Wait.",
    },
    {
        "id": "MT-A-009",
        "title": "Argumentos inválidos y ayuda",
        "objective": "Validar la interfaz pública y los códigos de salida básicos.",
        "preconditions": "Build Debug aprobado.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli help; confirmar $LASTEXITCODE igual a 0."),
            ("3", "& $cli; confirmar $LASTEXITCODE igual a 2."),
            ("4", "& $cli break '.\\source\\Scripts\\TestScript.json' abc; confirmar $LASTEXITCODE igual a 2."),
        ],
        "expected": "help devuelve 0; ausencia de argumentos y breakpoint inválido devuelven 2; ninguno inicia hardware.",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)
    for cell in row.cells:
        set_cell_shading(cell, TABLE_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(INK)


def set_run_font(run, size=11, color=INK, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_table(document: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]], widths: list[int]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = value
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_test_case(document: Document, test: dict) -> None:
    document.add_page_break()
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(f"{test['id']} | {test['title']}")

    add_table(
        document,
        ("Campo", "Detalle"),
        [
            ("Objetivo", test["objective"]),
            ("Precondiciones", test["preconditions"]),
        ],
        [1700, 7660],
    )

    document.add_paragraph("Procedimiento", style="Heading 2")
    add_table(document, ("Paso", "Acción exacta"), test["steps"], [850, 8510])

    document.add_paragraph("Resultado esperado", style="Heading 2")
    add_table(document, ("Criterio", "Resultado"), [("Aceptación", test["expected"])], [1700, 7660])

    document.add_paragraph("Registro de ejecución", style="Heading 2")
    add_table(
        document,
        ("Campo", "Dato / evidencia"),
        [
            ("Fecha y hora", ""),
            ("Ejecutado por", ""),
            ("Commit / rama", ""),
            ("Configuración", ""),
            ("Resultado real", ""),
            ("Código de salida", ""),
            ("Veredicto", "[ ] PASSED    [ ] FAILED    [ ] BLOCKED"),
            ("Incidencia asociada", ""),
        ],
        [2100, 7260],
    )

    document.add_paragraph("Evidencia", style="Heading 2")
    add_table(
        document,
        ("Tipo", "Archivo, captura o referencia"),
        [
            ("Captura 1", ""),
            ("Captura 2", ""),
            ("Log / reporte", ""),
            ("Observaciones", ""),
        ],
        [2100, 7260],
    )


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    document.core_properties.title = "ARTestCLI - Reporte de pruebas manuales - Etapa A"
    document.core_properties.subject = "Base segura de ARTestCLI"
    document.core_properties.author = "ARTest Engineering"
    document.core_properties.keywords = "ARTestCLI, QA, regression, Stage A"

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("ARTestCLI | Etapa A - Base segura")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Control de calidad | Página ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(footer)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("REPORTE DE PRUEBAS MANUALES")
    set_run_font(run, size=23, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("ARTestCLI - Etapa A: Base segura")
    set_run_font(run, size=14, color=MUTED)

    add_table(
        document,
        ("Control", "Valor"),
        [
            ("Documento", "ARTestCLI-QA-MAN-A"),
            ("Versión", "1.0"),
            ("Fecha de emisión", date.today().isoformat()),
            ("Configuración objetivo", "Visual Studio 18 Insiders | v145 | C++20 | x64"),
            ("Responsable", ""),
            ("Commit evaluado", ""),
            ("Veredicto general", "[ ] PASSED    [ ] FAILED    [ ] BLOCKED"),
        ],
        [2300, 7060],
    )

    document.add_paragraph("Propósito", style="Heading 1")
    document.add_paragraph(
        "Registrar evidencia reproducible de la línea base segura de ARTestCLI antes "
        "de extraer el motor a una DLL o introducir el contrato de plugins."
    )

    document.add_paragraph("Reglas de ejecución", style="Heading 1")
    add_table(
        document,
        ("Regla", "Aplicación"),
        [
            ("Trazabilidad", "Anotar commit, configuración, fecha, ejecutor y código de salida."),
            ("Evidencia", "Adjuntar captura legible y, cuando aplique, reporte XML/HTML."),
            ("Veredicto", "Marcar una sola opción. FAILED requiere una incidencia asociada."),
            ("Secuencia", "Ejecutar MT-A-001 a MT-A-009 en orden."),
            ("Cierre", "El veredicto general sólo puede ser PASSED si todos los casos son PASSED."),
        ],
        [2100, 7260],
    )

    document.add_paragraph("Resumen de casos", style="Heading 1")
    summary_rows = [(case["id"], case["title"], "") for case in TEST_CASES]
    add_table(document, ("ID", "Caso", "Veredicto"), summary_rows, [1200, 6360, 1800])

    for test in TEST_CASES:
        add_test_case(document, test)

    document.add_page_break()
    document.add_paragraph("Aprobación y cierre", style="Heading 1")
    add_table(
        document,
        ("Rol", "Nombre / firma / fecha"),
        [
            ("Ejecutor QA", ""),
            ("Revisor técnico", ""),
            ("Aprobación de fase", ""),
        ],
        [2500, 6860],
    )
    document.add_paragraph("Conclusión", style="Heading 2")
    add_table(
        document,
        ("Resultado", "Justificación y riesgos pendientes"),
        [("[ ] PASSED    [ ] FAILED    [ ] BLOCKED", "")],
        [3300, 6060],
    )
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
