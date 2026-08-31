from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = (
    REPOSITORY_ROOT
    / "quality"
    / "manual-tests"
    / "stage-b"
    / "ARTestCLI_Manual_Test_Report_Engine_Core_v1.0.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5D6B79"
TABLE_FILL = "E8EEF5"
WHITE = "FFFFFF"


TEST_CASES = [
    {
        "id": "MT-B-001",
        "title": "Compilacion y regresion Debug x64",
        "objective": "Validar la solucion de tres proyectos, el enlace del Core y la suite automatizada Debug.",
        "preconditions": "Visual Studio 18 Insiders instalado en D:\\Program Files\\Microsoft Visual Studio\\18\\Insiders.",
        "steps": [
            ("1", "Abrir PowerShell en D:\\GitHub\\main\\ARTestCLI."),
            ("2", "Ejecutar: .\\scripts\\build.ps1 -Configuration Debug -Platform x64"),
            ("3", "Esperar a que terminen MSBuild, Google Test y la validacion del reporte HTML."),
            ("4", "Ejecutar: $LASTEXITCODE y adjuntar captura de la salida final."),
        ],
        "expected": "Codigo 0; ARTestEngine.Core.lib, ARTestCLI.exe y ARTestCLI.UnitTests.exe generados; 25/25 pruebas PASSED; reporte consistente con Failed 0.",
    },
    {
        "id": "MT-B-002",
        "title": "Compilacion y regresion Release x64",
        "objective": "Confirmar que la arquitectura extraida es estable con optimizaciones Release.",
        "preconditions": "MT-B-001 aprobado.",
        "steps": [
            ("1", "Ejecutar: .\\scripts\\build.ps1 -Configuration Release -Platform x64"),
            ("2", "Abrir artifacts\\test-results\\x64\\Release\\ARTestCLI.UnitTests.html."),
            ("3", "Registrar $LASTEXITCODE y adjuntar evidencia del resumen."),
        ],
        "expected": "Codigo 0; 25/25 pruebas PASSED; Overall PASSED, Failed 0 y reporte HTML consistente.",
    },
    {
        "id": "MT-B-003",
        "title": "Limites de proyecto en Visual Studio",
        "objective": "Verificar que Core, CLI y pruebas son proyectos separados con dependencias unidireccionales.",
        "preconditions": "Abrir source\\ARTestCLI.sln en Visual Studio Insiders; seleccionar x64 Debug.",
        "steps": [
            ("1", "Confirmar que Solution Explorer muestra ARTestEngine.Core, ARTestCLI y ARTestCLI.UnitTests."),
            ("2", "Abrir Project Dependencies y confirmar que ARTestCLI depende de ARTestEngine.Core."),
            ("3", "Confirmar que ARTestCLI.UnitTests depende de ARTestEngine.Core."),
            ("4", "Compilar con Build > Build Solution y ejecutar Test > Run All Tests."),
        ],
        "expected": "Los tres proyectos compilan; no hay dependencia Core -> CLI; Test Explorer descubre 25 pruebas en 7 suites y todas pasan.",
    },
    {
        "id": "MT-B-004",
        "title": "Compilacion offline sin inicializar instrumentos",
        "objective": "Comprobar la separacion parser/compilador/executor y que compile no abre recursos.",
        "preconditions": "Build Debug aprobado.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli compile '.\\source\\Scripts\\TestScript.json'"),
            ("3", "Registrar $LASTEXITCODE y conservar la salida completa."),
        ],
        "expected": "Codigo 0 y texto 'Valid script. No instruments were initialized.'; no aparecen Initialize, Executing step, operaciones ni Shutdown.",
    },
    {
        "id": "MT-B-005",
        "title": "Ejecucion con instrumentos falsos y eventos",
        "objective": "Validar que el host presenta eventos del Core y que los fakes ejecutan la secuencia sin hardware.",
        "preconditions": "Build Debug aprobado; no conectar hardware fisico.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\source\\Scripts\\TestScript.json'"),
            ("3", "Registrar $LASTEXITCODE y adjuntar toda la salida."),
        ],
        "expected": "Codigo 0; inicializa CAN1 y PS1; publica operaciones PowerSupply/CAN; ejecuta pasos 1 a 4; finaliza PASSED y apaga ambos instrumentos.",
    },
    {
        "id": "MT-B-006",
        "title": "Rollback ante falla parcial de inicializacion",
        "objective": "Demostrar que una falla posterior libera los instrumentos inicializados previamente.",
        "preconditions": "Fixture quality\\manual-tests\\stage-b\\data\\initialization-rollback.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-b\\data\\initialization-rollback.json'"),
            ("3", "Ejecutar: $LASTEXITCODE y revisar el orden completo de mensajes."),
        ],
        "expected": "Codigo 4; A_CAN se inicializa; B_PS reporta POWER_SUPPLY_INITIALIZATION_FORCED_FAILURE; A_CAN ejecuta Shutdown; no se ejecuta el paso 1.",
    },
    {
        "id": "MT-B-007",
        "title": "Rechazo atomico de comando desconocido",
        "objective": "Confirmar que el compilador semantico no entrega una secuencia parcial.",
        "preconditions": "Fixture quality\\manual-tests\\stage-a\\data\\unknown-command.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-a\\data\\unknown-command.json'"),
            ("3", "Registrar $LASTEXITCODE y revisar que no existan eventos de inicializacion o pasos."),
        ],
        "expected": "Codigo 3 y diagnostico COMMAND_TYPE_UNKNOWN; ningun instrumento se inicializa y ningun paso se ejecuta.",
    },
    {
        "id": "MT-B-008",
        "title": "Control interactivo debug",
        "objective": "Validar que la interaccion de consola esta fuera del Core mediante IExecutionControl.",
        "preconditions": "Build Debug aprobado.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli debug '.\\source\\Scripts\\TestScript.json'"),
            ("3", "En la primera pausa escribir n y Enter; en la segunda escribir c y Enter."),
            ("4", "Confirmar que el resto de la secuencia continua y registrar $LASTEXITCODE."),
        ],
        "expected": "Pausa antes de los pasos 1 y 2; n conserva paso a paso; c continua; ejecucion PASSED y codigo 0.",
    },
    {
        "id": "MT-B-009",
        "title": "Breakpoint por indice de comando",
        "objective": "Comprobar el adaptador de breakpoints del host sin modificar el executor.",
        "preconditions": "Build Debug aprobado; indices base cero.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli break '.\\source\\Scripts\\TestScript.json' 1"),
            ("3", "Confirmar que el paso 1 se ejecuta sin pausa y que aparece una pausa antes del stepId 2."),
            ("4", "Escribir c y Enter; registrar $LASTEXITCODE."),
        ],
        "expected": "Existe una sola pausa en el indice 1 (stepId 2); despues de c la ejecucion termina PASSED con codigo 0.",
    },
    {
        "id": "MT-B-010",
        "title": "Aislamiento de consola en ARTestEngine.Core",
        "objective": "Verificar que el motor reusable no contiene dependencias directas de entrada/salida de consola.",
        "preconditions": "ripgrep (rg) disponible en la terminal del repositorio.",
        "steps": [
            ("1", "Ejecutar: rg -n 'std::(cout|cerr|cin)|#include\\s*<iostream>' source\\ARTestEngine.Core"),
            ("2", "Confirmar que el comando no imprime coincidencias."),
            ("3", "Ejecutar: rg -n 'REGISTER_COMMAND|REGISTER_INSTRUMENT' source"),
            ("4", "Confirmar que tampoco se imprimen coincidencias."),
        ],
        "expected": "Ambas busquedas regresan sin coincidencias: Core no usa consola y no existe autorregistro estatico por macros.",
    },
    {
        "id": "MT-B-011",
        "title": "Contrato de argumentos y codigos de salida",
        "objective": "Confirmar que el host delgado conserva el contrato de automatizacion de la Etapa A.",
        "preconditions": "Build Debug aprobado.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli help; registrar $LASTEXITCODE."),
            ("3", "& $cli; registrar $LASTEXITCODE."),
            ("4", "& $cli break '.\\source\\Scripts\\TestScript.json' abc; registrar $LASTEXITCODE."),
            ("5", "& $cli compile '.\\source\\Scripts\\missing.json'; registrar $LASTEXITCODE."),
        ],
        "expected": "help devuelve 0; sin argumentos y breakpoint invalido devuelven 2; archivo ausente devuelve 3; ningun caso invalido inicializa instrumentos.",
    },
]


def set_run_font(run, size=11, color=INK, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")


def shade_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), TABLE_FILL)
        cell._tc.get_or_add_tcPr().append(shading)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_geometry(table, widths)
    shade_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.25, bold=row_index == 0)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def configure_styles(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def add_test_case(document, case) -> None:
    document.add_page_break()
    document.add_paragraph(f"{case['id']} | {case['title']}", style="Heading 1")
    add_table(
        document,
        ("Campo", "Detalle"),
        [("Objetivo", case["objective"]), ("Precondiciones", case["preconditions"])],
        [1700, 7660],
    )
    document.add_paragraph("Procedimiento", style="Heading 2")
    add_table(document, ("Paso", "Accion exacta"), case["steps"], [850, 8510])
    document.add_paragraph("Resultado esperado", style="Heading 2")
    add_table(document, ("Criterio", "Resultado"), [("Aceptacion", case["expected"])], [1700, 7660])
    document.add_paragraph("Registro y evidencia", style="Heading 2")
    add_table(
        document,
        ("Campo", "Dato / evidencia"),
        [
            ("Fecha / ejecutor", ""),
            ("Commit / rama", ""),
            ("Resultado real", ""),
            ("Codigo de salida", ""),
            ("Veredicto", "[ ] PASSED    [ ] FAILED    [ ] BLOCKED"),
            ("Capturas / reporte", ""),
            ("Incidencia / notas", ""),
        ],
        [2100, 7260],
    )


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    document.core_properties.title = "ARTestCLI - Reporte de pruebas manuales - Etapa B"
    document.core_properties.subject = "Extraccion de ARTestEngine.Core"
    document.core_properties.author = "ARTest Engineering"
    document.core_properties.keywords = "ARTestCLI, ARTestEngine.Core, QA, regression, Stage B"

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("ARTestCLI | Etapa B - Engine Core"), size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Control de calidad | Pagina "), size=9, color=MUTED)
    add_page_field(footer)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("REPORTE DE PRUEBAS MANUALES"), size=23, color=INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("ARTestCLI - Etapa B: Extraccion de ARTestEngine.Core"), size=14, color=MUTED)

    add_table(
        document,
        ("Control", "Valor"),
        [
            ("Documento", "ARTestCLI-QA-MAN-B"),
            ("Version", "1.0"),
            ("Fecha de emision", date.today().isoformat()),
            ("Configuracion", "Visual Studio 18 Insiders | v145 | C++20 | x64"),
            ("Responsable", ""),
            ("Commit evaluado", ""),
            ("Veredicto general", "[ ] PASSED    [ ] FAILED    [ ] BLOCKED"),
        ],
        [2300, 7060],
    )
    document.add_paragraph("Proposito", style="Heading 1")
    document.add_paragraph(
        "Registrar evidencia reproducible de que el motor fue separado del host CLI, "
        "que parser, compilador y executor conservan el comportamiento y que la nueva "
        "arquitectura opera sin hardware fisico."
    )
    document.add_paragraph("Criterio de cierre", style="Heading 1")
    add_table(
        document,
        ("Regla", "Aplicacion"),
        [
            ("Trazabilidad", "Anotar commit, rama, configuracion, fecha y ejecutor."),
            ("Evidencia", "Adjuntar capturas legibles y los reportes HTML cuando correspondan."),
            ("Fallas", "Todo FAILED requiere incidencia, resultado real y evidencia."),
            ("Cierre", "El veredicto general solo puede ser PASSED si MT-B-001 a MT-B-011 son PASSED."),
        ],
        [2100, 7260],
    )
    document.add_paragraph("Resumen de casos", style="Heading 1")
    add_table(
        document,
        ("ID", "Caso", "Veredicto"),
        [(case["id"], case["title"], "") for case in TEST_CASES],
        [1200, 6360, 1800],
    )

    for case in TEST_CASES:
        add_test_case(document, case)

    document.add_page_break()
    document.add_paragraph("Aprobacion y cierre", style="Heading 1")
    add_table(
        document,
        ("Rol", "Nombre / firma / fecha"),
        [("Ejecutor QA", ""), ("Revisor tecnico", ""), ("Aprobacion de etapa", "")],
        [2500, 6860],
    )
    document.add_paragraph("Conclusion", style="Heading 2")
    add_table(
        document,
        ("Resultado", "Justificacion y riesgos pendientes"),
        [("[ ] PASSED    [ ] FAILED    [ ] BLOCKED", "")],
        [3300, 6060],
    )
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
