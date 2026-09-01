from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = (
    REPOSITORY_ROOT
    / "quality"
    / "manual-tests"
    / "stage-c"
    / "ARTestCLI_Manual_Test_Report_Robust_Execution_v1.0.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5D6B79"
TABLE_FILL = "E8EEF5"


TEST_CASES = [
    {
        "id": "MT-C-001",
        "title": "Compilacion y regresion Debug x64",
        "objective": "Validar la integracion completa de la Etapa C en configuracion Debug.",
        "preconditions": "Visual Studio 18 Insiders instalado en D:\\Program Files\\Microsoft Visual Studio\\18\\Insiders.",
        "steps": [
            ("1", "Abrir PowerShell en D:\\GitHub\\main\\ARTestCLI."),
            ("2", "Ejecutar: .\\scripts\\build.ps1 -Configuration Debug -Platform x64"),
            ("3", "Esperar a que terminen MSBuild, Google Test y la validacion HTML."),
            ("4", "Ejecutar: $LASTEXITCODE"),
            ("5", "Abrir artifacts\\test-results\\x64\\Debug\\ARTestCLI.UnitTests.html y adjuntar el resumen."),
        ],
        "expected": "Codigo 0; 41/41 pruebas PASSED en 11 suites; Overall PASSED y Failed 0.",
    },
    {
        "id": "MT-C-002",
        "title": "Compilacion y regresion Release x64",
        "objective": "Detectar diferencias por optimizacion y configuracion Release.",
        "preconditions": "MT-C-001 aprobado.",
        "steps": [
            ("1", "Ejecutar: .\\scripts\\build.ps1 -Configuration Release -Platform x64"),
            ("2", "Ejecutar: $LASTEXITCODE"),
            ("3", "Abrir artifacts\\test-results\\x64\\Release\\ARTestCLI.UnitTests.html y adjuntar el resumen."),
        ],
        "expected": "Codigo 0; 41/41 pruebas PASSED en 11 suites; Overall PASSED y Failed 0.",
    },
    {
        "id": "MT-C-003",
        "title": "Ciclo de estados y resumen exitoso",
        "objective": "Comprobar el ciclo normal de ExecutionSession y el reporte general.",
        "preconditions": "Build Debug aprobado.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\source\\Scripts\\TestScript.json'"),
            ("3", "$code = $LASTEXITCODE; $code"),
            ("4", "Conservar la salida completa desde INITIALIZING hasta el resumen."),
        ],
        "expected": "Codigo 0; estados INITIALIZING, RUNNING, CLEANING_UP y COMPLETED en ese orden; todos los steps PASSED; skipped=0; se observa Shutdown de cada instrumento.",
    },
    {
        "id": "MT-C-004",
        "title": "Retry hasta recuperar un fallo transitorio",
        "objective": "Validar maxAttempts, retryDelayMs y el registro por intento.",
        "preconditions": "Fixture quality\\manual-tests\\stage-c\\data\\retry-success.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-c\\data\\retry-success.json'"),
            ("3", "$code = $LASTEXITCODE; $code"),
            ("4", "Adjuntar salida con los tres intentos y el resumen."),
        ],
        "expected": "Codigo 0; el step 1 falla dos veces, programa dos retries y pasa en el intento 3; attempts=4 en el resumen total; estado final COMPLETED.",
    },
    {
        "id": "MT-C-005",
        "title": "Timeout por intento y detencion de secuencia",
        "objective": "Confirmar que un timeout cooperativo interrumpe Wait y omite pasos posteriores.",
        "preconditions": "Fixture quality\\manual-tests\\stage-c\\data\\timeout-stop.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "$started = Get-Date"),
            ("3", "& $cli run '.\\quality\\manual-tests\\stage-c\\data\\timeout-stop.json'"),
            ("4", "$code = $LASTEXITCODE; $elapsed = (Get-Date) - $started"),
            ("5", "\"Exit=$code ElapsedMs=$([math]::Round($elapsed.TotalMilliseconds))\""),
        ],
        "expected": "Codigo 5; step 1 TIMED_OUT; estado final TIMED_OUT; executed=1, timedOut=1 y skipped=1; duracion claramente menor que los 5000 ms solicitados por Wait.",
    },
    {
        "id": "MT-C-006",
        "title": "Politica continue conserva el fallo general",
        "objective": "Verificar que onFailure=continue ejecuta el siguiente step sin ocultar el error previo.",
        "preconditions": "Fixture quality\\manual-tests\\stage-c\\data\\continue-on-failure.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-c\\data\\continue-on-failure.json'"),
            ("3", "$code = $LASTEXITCODE; $code"),
            ("4", "Revisar los veredictos de ambos steps y el resumen."),
        ],
        "expected": "Codigo 5; step 1 ERROR; step 2 PASSED; executed=2, passed=1, errors=1 y skipped=0; estado final FAILED y veredicto general ERROR.",
    },
    {
        "id": "MT-C-007",
        "title": "Cancelacion real con Ctrl+C y cleanup",
        "objective": "Validar el adaptador de consola, la cancelacion cooperativa y el cleanup garantizado.",
        "preconditions": "Fixture quality\\manual-tests\\stage-c\\data\\cancel-cleanup.json disponible; ejecutar en consola interactiva.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-c\\data\\cancel-cleanup.json'"),
            ("3", "Esperar hasta que aparezca: Executing step 2: Time.WaitMs"),
            ("4", "Presionar Ctrl+C una sola vez y esperar el retorno del proceso."),
            ("5", "$code = $LASTEXITCODE; $code"),
            ("6", "Adjuntar la salida desde CANCELLING hasta el resumen final."),
        ],
        "expected": "Codigo 5; estados CANCELLING, CLEANING_UP y CANCELLED; step 2 CANCELLED; aparece PowerSupply Shutdown; step 3 no se ejecuta y skipped=1.",
    },
    {
        "id": "MT-C-008",
        "title": "Fallo de cleanup domina el veredicto",
        "objective": "Confirmar que una secuencia aprobada no se reporta PASSED si el instrumento no puede apagarse.",
        "preconditions": "Fixture quality\\manual-tests\\stage-c\\data\\cleanup-failure.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli run '.\\quality\\manual-tests\\stage-c\\data\\cleanup-failure.json'"),
            ("3", "$code = $LASTEXITCODE; $code"),
            ("4", "Adjuntar el diagnostico y el resumen final."),
        ],
        "expected": "Codigo 5; el step pasa; aparece POWER_SUPPLY_SHUTDOWN_FORCED_FAILURE; estado final FAILED y veredicto general ERROR.",
    },
    {
        "id": "MT-C-009",
        "title": "Rechazo de politica insegura",
        "objective": "Comprobar validacion semantica y rechazo atomico antes de inicializar instrumentos.",
        "preconditions": "Fixture quality\\manual-tests\\stage-c\\data\\invalid-policy.json disponible.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli compile '.\\quality\\manual-tests\\stage-c\\data\\invalid-policy.json'"),
            ("3", "$code = $LASTEXITCODE; $code"),
            ("4", "Confirmar que no aparecen INITIALIZING ni eventos de instrumentos."),
        ],
        "expected": "Codigo 3 y diagnostico COMMAND_POLICY_ATTEMPTS_INVALID; no se inicializan instrumentos ni se ejecutan steps.",
    },
    {
        "id": "MT-C-010",
        "title": "Compatibilidad con scripts version 1 sin policy",
        "objective": "Demostrar que policy es opcional y no rompe documentos existentes.",
        "preconditions": "source\\Scripts\\TestScript.json no contiene objetos policy.",
        "steps": [
            ("1", "$cli = '.\\artifacts\\bin\\x64\\Debug\\ARTestCLI.exe'"),
            ("2", "& $cli compile '.\\source\\Scripts\\TestScript.json'; $compileCode = $LASTEXITCODE"),
            ("3", "& $cli run '.\\source\\Scripts\\TestScript.json'; $runCode = $LASTEXITCODE"),
            ("4", "\"Compile=$compileCode Run=$runCode\""),
        ],
        "expected": "Compile=0 y Run=0; un solo intento por step; sin retries ni timeout; ejecucion final PASSED.",
    },
    {
        "id": "MT-C-011",
        "title": "Descubrimiento en Visual Studio Test Explorer",
        "objective": "Verificar que las pruebas de Etapa C estan integradas en la solucion.",
        "preconditions": "Abrir source\\ARTestCLI.sln en Visual Studio Insiders; seleccionar x64 Debug.",
        "steps": [
            ("1", "Abrir Test > Test Explorer."),
            ("2", "Seleccionar Build > Build Solution."),
            ("3", "Confirmar que se descubren 41 pruebas en 11 suites."),
            ("4", "Seleccionar Run All Tests y adjuntar el resumen de Test Explorer."),
        ],
        "expected": "41 pruebas descubiertas y 41 PASSED; no hay FAILED, SKIPPED ni pruebas sin ejecutar.",
    },
    {
        "id": "MT-C-012",
        "title": "Consistencia del reporte automatizado",
        "objective": "Confirmar que el resumen HTML coincide con cada test case individual.",
        "preconditions": "MT-C-001 y MT-C-002 aprobados.",
        "steps": [
            ("1", "Abrir los reportes HTML Debug y Release."),
            ("2", "Confirmar Overall PASSED, Total 41, Passed 41 y Failed 0 en ambos."),
            ("3", "Recorrer la columna Status y confirmar que cada caso indica PASSED."),
            ("4", "Adjuntar capturas del resumen y de la tabla de casos."),
        ],
        "expected": "Los contadores agregados y los 41 veredictos individuales coinciden en Debug y Release; no existe contradiccion PASSED/FAILED.",
    },
]


def set_run_font(run, size=11, color=INK, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")


def shade_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), TABLE_FILL)
        cell._tc.get_or_add_tcPr().append(shading)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_geometry(table, widths)
    shade_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.25, bold=row_index == 0)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def configure_styles(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def add_test_case(document, case) -> None:
    document.add_page_break()
    document.add_paragraph(f"{case['id']} | {case['title']}", style="Heading 1")
    add_table(
        document,
        ("Campo", "Detalle"),
        [("Objetivo", case["objective"]), ("Precondiciones", case["preconditions"])],
        [1700, 7660],
    )
    document.add_paragraph("Procedimiento", style="Heading 2")
    add_table(document, ("Paso", "Accion exacta"), case["steps"], [850, 8510])
    document.add_paragraph("Resultado esperado", style="Heading 2")
    add_table(document, ("Criterio", "Resultado"), [("Aceptacion", case["expected"])], [1700, 7660])
    document.add_paragraph("Registro y evidencia", style="Heading 2")
    add_table(
        document,
        ("Campo", "Dato / evidencia"),
        [
            ("Fecha / ejecutor", ""),
            ("Commit / rama", ""),
            ("Resultado real", ""),
            ("Codigo de salida", ""),
            ("Veredicto", "( ) PASSED    ( ) FAILED    ( ) BLOCKED"),
            ("Capturas / reporte", ""),
            ("Incidencia / notas", ""),
        ],
        [2100, 7260],
    )


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    document.core_properties.title = "ARTestCLI - Reporte de pruebas manuales - Etapa C"
    document.core_properties.subject = "Ejecucion robusta"
    document.core_properties.author = "ARTest Engineering"
    document.core_properties.keywords = "ARTestCLI, execution, cancellation, timeout, retry, cleanup, QA"

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("ARTestCLI | Etapa C - Ejecucion robusta"), size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Control de calidad | Pagina "), size=9, color=MUTED)
    add_page_field(footer)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("REPORTE DE PRUEBAS MANUALES"), size=23, color=INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("ARTestCLI - Etapa C: Ejecucion robusta"), size=14, color=MUTED)

    add_table(
        document,
        ("Control", "Valor"),
        [
            ("Documento", "ARTestCLI-QA-MAN-C"),
            ("Version", "1.0"),
            ("Fecha de emision", date.today().isoformat()),
            ("Configuracion", "Visual Studio 18 Insiders | v145 | C++20 | x64"),
            ("Responsable", ""),
            ("Commit evaluado", ""),
            ("Veredicto general", "( ) PASSED    ( ) FAILED    ( ) BLOCKED"),
        ],
        [2300, 7060],
    )
    document.add_paragraph("Proposito", style="Heading 1")
    document.add_paragraph(
        "Registrar evidencia reproducible de la maquina de estados, ejecucion asincrona, "
        "cancelacion, timeout, retry, politicas de fallo, cleanup garantizado y reportes "
        "por step y por ejecucion incorporados en la Etapa C."
    )
    document.add_paragraph("Criterio de cierre", style="Heading 1")
    add_table(
        document,
        ("Regla", "Aplicacion"),
        [
            ("Trazabilidad", "Anotar commit, rama, configuracion, fecha y ejecutor."),
            ("Evidencia", "Adjuntar capturas legibles y reportes HTML cuando correspondan."),
            ("Seguridad", "Cancelacion, timeout y fallos deben pasar por CLEANING_UP."),
            ("Fallas", "Todo FAILED requiere incidencia, resultado real y evidencia."),
            ("Cierre", "El veredicto general solo puede ser PASSED si MT-C-001 a MT-C-012 son PASSED."),
        ],
        [2100, 7260],
    )
    document.add_paragraph("Resumen de casos", style="Heading 1")
    add_table(
        document,
        ("ID", "Caso", "Veredicto"),
        [(case["id"], case["title"], "") for case in TEST_CASES],
        [1200, 6360, 1800],
    )

    for case in TEST_CASES:
        add_test_case(document, case)

    document.add_page_break()
    document.add_paragraph("Aprobacion y cierre", style="Heading 1")
    add_table(
        document,
        ("Rol", "Nombre / firma / fecha"),
        [("Ejecutor QA", ""), ("Revisor tecnico", ""), ("Aprobacion de etapa", "")],
        [2500, 6860],
    )
    document.add_paragraph("Conclusion", style="Heading 2")
    add_table(
        document,
        ("Resultado", "Justificacion y riesgos pendientes"),
        [("( ) PASSED    ( ) FAILED    ( ) BLOCKED", "")],
        [3300, 6060],
    )
    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
