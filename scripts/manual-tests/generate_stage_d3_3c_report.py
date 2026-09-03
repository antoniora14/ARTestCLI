from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from generate_stage_d1_report import (
    BLUE,
    DARK_BLUE,
    NAVY,
    add_evidence_block,
    add_metadata,
    add_page_number,
    add_steps,
    configure_styles,
    set_cell_text,
    set_font,
    set_table_geometry,
    style_table,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = (
    ROOT
    / "quality"
    / "manual-tests"
    / "stage-d3.3c"
    / "ARTestCLI_Manual_Test_Report_SDK_Distribution_v1.0.docx"
)


# Compact reference guide; memo-style cover, matching the earlier ARTest QA reports.
# Named overrides: QA table text 9.5 pt / 1.15 spacing; title 24 pt navy.
import json
TEST_CASES = json.loads((ROOT / "quality/manual-tests/stage-d3.3c/cases.json").read_text(encoding="utf-8"))



def build(output: Path = OUTPUT):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_font(
        header.add_run("ARTest Quality | Stage D3.3-C Manual Validation"),
        size=9,
        color="666666",
    )
    add_page_number(section.footer.paragraphs[0])

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_font(kicker.add_run("QUALITY ASSURANCE"), size=10, bold=True, color=BLUE)
    title = document.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(5)
    set_font(title.add_run("ARTestCLI Manual Test Report"), size=24, bold=True, color=NAVY)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(
        subtitle.add_run("Stage D3.3-C - SDK Distribution and External Consumer Compatibility"),
        size=14,
        color=DARK_BLUE,
    )

    add_metadata(
        document,
        (
            ("Document ID", "ARTEST-QA-D3.3C-MANUAL-001"),
            ("Version", "1.0"),
            ("Status", "Pending manual execution"),
            ("Prepared for", "ARTest SDK distribution acceptance"),
            ("Tester", ""),
            ("Execution date", ""),
            ("Commit / branch", ""),
        ),
    )

    document.add_heading("Purpose and acceptance scope", level=1)
    document.add_paragraph(
        "This report captures repeatable manual evidence for Stage D3.3-C SDK "
        "distribution. It validates package structure and inventory, consumption "
        "from the extracted ZIP, an external /W4 /WX extension build, Engine "
        "activation and execution, cleanup, and the official Release regression."
    )
    acceptance = document.add_paragraph()
    acceptance.add_run("Acceptance rule. ").bold = True
    acceptance.add_run(
        "All three manual cases must be Passed. The Release automated report "
        "must show 139 passed, 0 failed, and 0 skipped. "
        "Manual evidence remains Pending until executed and recorded by the tester."
    )

    document.add_page_break()
    document.add_heading("Execution environment", level=1)
    add_metadata(
        document,
        (
            ("Operating system", "Windows:"),
            ("Visual Studio", "18 Insiders / v145"),
            ("Platform", "x64"),
            ("Configurations", "Release acceptance; Debug compatibility pre-check"),
            ("Repository", "D:\\GitHub\\main\\ARTestCLI"),
            ("PowerShell version", ""),
            ("Machine / tester", ""),
        ),
    )

    document.add_paragraph("Run all commands from D:\\GitHub\\main\\ARTestCLI. In the same PowerShell session define:")
    for command in (
        ".\\scripts\\package-sdk.ps1 -Configuration Release -Platform x64",
        ".\\scripts\\test-sdk-distribution.ps1 -Configuration Release -Platform x64",
        ".\\scripts\\build.ps1 -Configuration Release -Platform x64",
    ):
        p = document.add_paragraph(command)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(9)
    document.add_paragraph("No physical hardware is required. Generated SDK artifacts remain under artifacts and are not committed.")
    document.add_heading("Test summary", level=1)
    summary = document.add_table(rows=1, cols=4)
    for index, text in enumerate(
        ("ID", "Scenario", "Expected verdict", "Actual verdict")
    ):
        set_cell_text(summary.rows[0].cells[index], text, bold=True, color=NAVY)
    for case in TEST_CASES:
        cells = summary.add_row().cells
        set_cell_text(cells[0], case["id"], bold=True, color=NAVY)
        set_cell_text(cells[1], case["title"])
        set_cell_text(cells[2], "Passed", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], "Pending", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(summary, [1800, 4560, 1500, 1500])
    style_table(summary)

    document.add_page_break()
    for index, case in enumerate(TEST_CASES):
        # Use independent paragraphs instead of a multiline built-in heading.
        # Some Word versions apply a hanging first-line indent after an explicit
        # page break, which can move the case identifier outside the page margin.
        case_id = document.add_paragraph()
        case_id.paragraph_format.space_before = Pt(18)
        case_id.paragraph_format.space_after = Pt(3)
        case_id.paragraph_format.keep_with_next = True
        set_font(case_id.add_run(case["id"]), size=16, bold=True, color=BLUE)
        case_title = document.add_paragraph()
        case_title.paragraph_format.space_after = Pt(10)
        case_title.paragraph_format.keep_with_next = True
        set_font(case_title.add_run(case["title"]), size=16, bold=True, color=BLUE)
        add_metadata(
            document,
            (
                ("Objective", case["objective"]),
                ("Preconditions", case["preconditions"]),
                ("Expected verdict", "Passed"),
            ),
        )
        document.add_heading("Procedure and expected results", level=2)
        rows = [
            [number, step["action"], step["expected"]]
            for number, step in enumerate(case["steps"], start=1)
        ]
        add_steps(document, rows)
        set_table_geometry(document.tables[-1], [600, 4800, 3960])
        document.add_heading("Execution evidence", level=2)
        add_evidence_block(document)
        if index != len(TEST_CASES) - 1:
            document.add_page_break()

    document.add_page_break()
    document.add_heading("Final acceptance", level=1)
    add_metadata(
        document,
        (
            ("Total manual cases", str(len(TEST_CASES))),
            ("Passed", ""),
            ("Failed", ""),
            ("Blocked / Not run", ""),
            ("SDK Debug compatibility", ""),
            ("Automated Release regression", ""),
            ("Overall Stage D3.3-C verdict", "Pending"),
        ),
    )
    document.add_heading("Approval", level=2)
    approval = document.add_table(rows=3, cols=3)
    for index, text in enumerate(("Role", "Name / Signature", "Date")):
        set_cell_text(approval.rows[0].cells[index], text, bold=True, color=NAVY)
    for row_index, role in enumerate(("Tester", "Technical reviewer"), start=1):
        set_cell_text(approval.rows[row_index].cells[0], role, bold=True)
        set_cell_text(approval.rows[row_index].cells[1], "")
        set_cell_text(approval.rows[row_index].cells[2], "")
    set_table_geometry(approval, [1900, 5060, 2400])
    style_table(approval)

    core = document.core_properties
    core.title = "ARTestCLI Manual Test Report - Stage D3.3-C"
    core.subject = "SDK distribution and external consumer compatibility acceptance"
    core.author = "ARTest Quality"
    core.keywords = "ARTestCLI, ARTestEngine, D3.3-C, SDK, distribution, external consumer"
    output.parent.mkdir(parents=True, exist_ok=True)
    if output.exists():
        raise FileExistsError("Preserve existing manual evidence: choose another output filename.")
    document.save(output)
    print(output)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Create a blank D3.3-C evidence report without overwriting an existing file.")
    parser.add_argument("--output", type=Path, default=OUTPUT)
    build(parser.parse_args().output)
