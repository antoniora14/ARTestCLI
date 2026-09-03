from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from generate_stage_d1_report import (
    BLUE,
    DARK_BLUE,
    NAVY,
    add_evidence_block,
    add_metadata,
    add_note,
    add_page_number,
    add_steps,
    configure_styles,
    set_cell_text,
    set_font,
    set_table_geometry,
    style_table,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = (
    ROOT
    / "quality"
    / "manual-tests"
    / "stage-d3.1"
    / "ARTestCLI_Manual_Test_Report_Production_Catalog_v1.0.docx"
)


TEST_CASES = [
    {
        "id": "MT-D3.1-001",
        "title": "Official Debug and Release regression",
        "objective": "Confirm the solution, ABI layouts, catalog security checks, CLI contracts, and validated reports pass in both supported configurations.",
        "preconditions": "Visual Studio 18 Insiders is installed at the documented path. Start in D:\\GitHub\\main\\ARTestCLI.",
        "steps": [
            (1, "Run .\\scripts\\build.ps1 -Configuration Debug -Platform x64.", "Build, ABI contract, thin-host verification, report checks, and all 72 Google Tests pass across 19 suites."),
            (2, "Open artifacts\\test-results\\x64\\Debug\\ARTestCLI.UnitTests.html.", "Overall is PASSED; Total and Passed are 72; Failed and Skipped are 0; every case verdict is consistent."),
            (3, "Run .\\scripts\\build.ps1 -Configuration Release -Platform x64.", "Release produces the same 72/72 passed baseline and a validated HTML report."),
            (4, "Open the Release HTML report and record both reports as evidence.", "Debug and Release configuration, platform, counts, and final verdict are visible."),
        ],
    },
    {
        "id": "MT-D3.1-002",
        "title": "Deterministic package list and packaged integrity",
        "objective": "Verify packaged manifests contain configuration-specific SHA-256 values and list presents a deterministic catalog summary.",
        "preconditions": "Release build completed successfully.",
        "steps": [
            (1, "Set $cli = '.\\artifacts\\bin\\x64\\Release\\ARTestCLI.exe' and $extensions = '.\\artifacts\\extensions\\x64\\Release'.", "Test-Path returns True for both paths."),
            (2, "Run & $cli extensions list $extensions; $LASTEXITCODE.", "Two VALID packages appear in extension-ID order, each reports sha256=verified, the summary says 2 package(s), catalog valid, and exit code is 0."),
            (3, "Run (Get-Content ($extensions + '\\ARTestCmdSample\\artest-extension.json') -Raw | ConvertFrom-Json).integrity.sha256.", "A lower-case 64-character hexadecimal SHA-256 value is displayed."),
        ],
    },
    {
        "id": "MT-D3.1-003",
        "title": "Side-effect-free catalog validation report",
        "objective": "Confirm validation returns the catalog v2 machine-readable verdict without activating extension code.",
        "preconditions": "$cli and $extensions are defined as in MT-D3.1-002.",
        "steps": [
            (1, "Run & $cli extensions validate $extensions; $LASTEXITCODE.", "The command returns 0 and emits one formatted JSON document."),
            (2, "Inspect schema, status, valid, generation, packages, and extensions.", "schema is artest.schema.extension-catalog.v2; status is validated; valid is true; generation is 0; two packages are valid and verified; extensions is empty."),
            (3, "Confirm no [Engine] extension-catalog activation line is printed.", "No native activation event appears because validate never loads DLLs."),
        ],
    },
    {
        "id": "MT-D3.1-004",
        "title": "Full native catalog doctor and activation",
        "objective": "Verify binary exports, ABI tables, descriptors, and registry conflicts are checked before activation.",
        "preconditions": "$cli and $extensions are defined.",
        "steps": [
            (1, "Run & $cli extensions doctor $extensions; $LASTEXITCODE.", "The command returns 0."),
            (2, "Inspect Engine event output.", "It states that the native extension catalog was validated and activated atomically."),
            (3, "Inspect the JSON snapshot.", "status is active; valid is true; generation is 1; two active extension manifests are present; diagnostics is empty."),
            (4, "Run & $cli extension-run '.\\source\\Scripts\\ExtensionScript.json' $extensions.", "The command-to-driver service flow passes and returns exit code 0."),
        ],
    },
    {
        "id": "MT-D3.1-005",
        "title": "Incompatible ABI rejection and dedicated exit code",
        "objective": "Confirm an incompatible manifest is rejected as catalog data and cannot reach native loading.",
        "preconditions": "$cli is defined. The versioned incompatible fixture exists.",
        "steps": [
            (1, "Set $bad = '.\\quality\\manual-tests\\stage-d1\\data\\incompatible'.", "The fixture root contains BadAbi\\artest-extension.json."),
            (2, "Run & $cli extensions validate $bad; $LASTEXITCODE.", "The report has valid=false, contains EXTENSION_RUNTIME_INCOMPATIBLE, and the process returns dedicated exit code 6."),
            (3, "Confirm the output does not contain an active status or activation event.", "The invalid catalog is not activated."),
        ],
    },
    {
        "id": "MT-D3.1-006",
        "title": "Tampered DLL integrity rejection",
        "objective": "Verify a changed binary cannot pass a manifest SHA-256 check.",
        "preconditions": "$cli and $extensions are defined. Use the commands exactly on a temporary copy.",
        "steps": [
            (1, "$tempCatalog = Join-Path $env:TEMP 'ARTest-D3-Integrity'; Remove-Item -LiteralPath $tempCatalog -Recurse -Force -ErrorAction SilentlyContinue; Copy-Item -LiteralPath $extensions -Destination $tempCatalog -Recurse.", "A disposable copy of the Release catalog is created."),
            (2, "$dll = Join-Path $tempCatalog 'ARTestCmdSample\\ARTestCmdSample.dll'; $stream = [IO.File]::OpenWrite($dll); try { $null = $stream.Seek(0,[IO.SeekOrigin]::End); $stream.WriteByte(0) } finally { $stream.Dispose() }.", "One byte is appended to the copied DLL; the original artifact remains unchanged."),
            (3, "Run & $cli extensions validate $tempCatalog; $LASTEXITCODE.", "valid is false, integrity is mismatch, EXTENSION_INTEGRITY_MISMATCH is present, and exit code is 6."),
            (4, "Remove-Item -LiteralPath $tempCatalog -Recurse -Force.", "The temporary catalog is removed."),
        ],
    },
    {
        "id": "MT-D3.1-007",
        "title": "Runtime entry path-containment rejection",
        "objective": "Verify a manifest cannot escape its package directory even when the target file exists.",
        "preconditions": "$cli and $extensions are defined. Use a new temporary copy.",
        "steps": [
            (1, "$tempCatalog = Join-Path $env:TEMP 'ARTest-D3-Path'; Remove-Item -LiteralPath $tempCatalog -Recurse -Force -ErrorAction SilentlyContinue; Copy-Item -LiteralPath $extensions -Destination $tempCatalog -Recurse.", "A disposable catalog copy is created."),
            (2, "Copy-Item -LiteralPath (Join-Path $tempCatalog 'ARTestCmdSample\\ARTestCmdSample.dll') -Destination (Join-Path $tempCatalog 'outside.dll'); $manifestPath = Join-Path $tempCatalog 'ARTestCmdSample\\artest-extension.json'; $manifest = Get-Content -LiteralPath $manifestPath -Raw | ConvertFrom-Json; $manifest.runtime.entry = '..\\outside.dll'; $manifest | ConvertTo-Json -Depth 32 | Set-Content -LiteralPath $manifestPath -Encoding utf8.", "The copied manifest points to an existing DLL outside its own package."),
            (3, "Run & $cli extensions validate $tempCatalog; $LASTEXITCODE.", "valid is false, EXTENSION_ENTRY_INVALID is present, and exit code is 6."),
            (4, "Remove-Item -LiteralPath $tempCatalog -Recurse -Force.", "The temporary catalog is removed."),
        ],
    },
    {
        "id": "MT-D3.1-008",
        "title": "ABI compatibility and failed-activation containment evidence",
        "objective": "Confirm API 0.1/0.2 callers remain protected and a rejected refresh does not poison the Engine instance.",
        "preconditions": "Debug or Release automated HTML report is open.",
        "steps": [
            (1, "Locate StageDEngineApiCompatibilityTests in the report.", "Both minor-one and minor-two negotiation tests are PASSED and verify no bytes beyond the negotiated table are overwritten."),
            (2, "Locate EngineFixture.FailedActivationDoesNotPoisonTheNextValidActivation.", "The test is PASSED; the same Engine rejects ABI 99, accepts the corrected manifest, and exposes active generation 1."),
            (3, "Locate the malformed, path-escape, duplicate-ID, integrity-mismatch, and validation-no-load tests.", "Every D3.1 security and failure-containment case is PASSED."),
        ],
    },
]


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_font(
        header.add_run("ARTest Quality | Stage D3.1 Manual Validation"),
        size=9,
        color="666666",
    )
    add_page_number(section.footer.paragraphs[0])

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_font(kicker.add_run("QUALITY ASSURANCE"), size=10, bold=True, color=BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_font(title.add_run("ARTestCLI Manual Test Report"), size=24, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(
        subtitle.add_run("Stage D3.1 - Production Extension Catalog"),
        size=14,
        color=DARK_BLUE,
    )

    add_metadata(
        document,
        (
            ("Document ID", "ARTEST-QA-D3.1-MANUAL-001"),
            ("Version", "1.0"),
            ("Status", "Pending manual execution"),
            ("Prepared for", "ARTest catalog D3.1 acceptance"),
            ("Tester", ""),
            ("Execution date", ""),
            ("Commit / branch", ""),
        ),
    )

    document.add_heading("Purpose and acceptance scope", level=1)
    document.add_paragraph(
        "This report captures repeatable manual evidence for the Stage D3.1 "
        "production extension catalog. It validates safe discovery, manifest "
        "and path rules, binary integrity, ABI inspection, transactional "
        "activation, compatibility, diagnostics, and CLI exit contracts."
    )
    add_note(
        document,
        "Acceptance rule",
        "All eight manual cases must be Passed, and the Debug and Release "
        "automated reports must each show 72 passed, 0 failed, and 0 skipped.",
    )

    document.add_page_break()
    document.add_heading("Execution environment", level=1)
    add_metadata(
        document,
        (
            ("Operating system", "Windows:"),
            ("Visual Studio", "18 Insiders / v145"),
            ("Platform", "x64"),
            ("Configurations", "Debug and Release"),
            ("Repository", "D:\\GitHub\\main\\ARTestCLI"),
            ("PowerShell version", ""),
            ("Machine / tester", ""),
        ),
    )

    document.add_heading("Test summary", level=1)
    summary = document.add_table(rows=1, cols=4)
    for index, text in enumerate(
        ("ID", "Scenario", "Expected verdict", "Actual verdict")
    ):
        set_cell_text(summary.rows[0].cells[index], text, bold=True, color=NAVY)
    for case in TEST_CASES:
        cells = summary.add_row().cells
        set_cell_text(cells[0], case["id"], bold=True, color=NAVY)
        set_cell_text(cells[1], case["title"])
        set_cell_text(cells[2], "Passed", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], "Pending", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(summary, [1250, 4910, 1600, 1600])
    style_table(summary)

    document.add_page_break()
    for index, case in enumerate(TEST_CASES):
        document.add_heading(f'{case["id"]} | {case["title"]}', level=1)
        add_metadata(
            document,
            (
                ("Objective", case["objective"]),
                ("Preconditions", case["preconditions"]),
                ("Expected verdict", "Passed"),
            ),
        )
        document.add_heading("Procedure and expected results", level=2)
        add_steps(document, case["steps"])
        document.add_heading("Execution evidence", level=2)
        add_evidence_block(document)
        if index != len(TEST_CASES) - 1:
            document.add_page_break()

    document.add_page_break()
    document.add_heading("Final acceptance", level=1)
    add_metadata(
        document,
        (
            ("Total manual cases", str(len(TEST_CASES))),
            ("Passed", ""),
            ("Failed", ""),
            ("Blocked / Not run", ""),
            ("Automated Debug verdict", ""),
            ("Automated Release verdict", ""),
            ("Overall Stage D3.1 verdict", "Pending"),
        ),
    )
    document.add_heading("Approval", level=2)
    approval = document.add_table(rows=3, cols=3)
    for index, text in enumerate(("Role", "Name / Signature", "Date")):
        set_cell_text(approval.rows[0].cells[index], text, bold=True, color=NAVY)
    for row_index, role in enumerate(("Tester", "Technical reviewer"), start=1):
        set_cell_text(approval.rows[row_index].cells[0], role, bold=True)
        set_cell_text(approval.rows[row_index].cells[1], "")
        set_cell_text(approval.rows[row_index].cells[2], "")
    set_table_geometry(approval, [1900, 5060, 2400])
    style_table(approval)

    core = document.core_properties
    core.title = "ARTestCLI Manual Test Report - Stage D3.1"
    core.subject = "Production extension catalog acceptance"
    core.author = "ARTest Quality"
    core.keywords = "ARTestCLI, ARTestEngine, D3.1, extension catalog, integrity"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
