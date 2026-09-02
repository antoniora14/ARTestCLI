from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from generate_stage_d1_report import (
    BLUE,
    DARK_BLUE,
    NAVY,
    add_evidence_block,
    add_metadata,
    add_note,
    add_page_number,
    add_steps,
    configure_styles,
    set_cell_text,
    set_font,
    set_table_geometry,
    style_table,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = (
    ROOT
    / "quality"
    / "manual-tests"
    / "stage-d2"
    / "ARTestCLI_Manual_Test_Report_Thin_Host_Migration_v1.0.docx"
)


TEST_CASES = [
    {
        "id": "MT-D2-001",
        "title": "Official Debug and Release regression workflow",
        "objective": "Confirm the complete solution, ABI layout, thin-host rule, Google Test suite, and validated reports pass in both supported configurations.",
        "preconditions": "Visual Studio 18 Insiders is installed at the documented path. Run from the repository root.",
        "steps": [
            (1, "Run .\\scripts\\build.ps1 -Configuration Debug -Platform x64.", "Build succeeds; thin-host verification passes; ABI contract returns 0; Google Test reports 61 passed and 0 failed."),
            (2, "Open artifacts\\test-results\\x64\\Debug\\ARTestCLI.UnitTests.html.", "Overall is PASSED; Total and Passed are 61; Failed and Skipped are 0; every test-case verdict is consistent."),
            (3, "Run .\\scripts\\build.ps1 -Configuration Release -Platform x64.", "Release produces the same 61/61 passed baseline and a validated HTML report."),
            (4, "Record both HTML reports and console output as evidence.", "Evidence identifies configuration, platform, suite count, test count, and final verdict."),
        ],
    },
    {
        "id": "MT-D2-002",
        "title": "Thin-host architectural boundary enforcement",
        "objective": "Verify ARTestCLI compiles only against ARTest.SDK and ARTestEngine.dll, never ARTestEngine.Core.",
        "preconditions": "Repository is on the D2 implementation under test.",
        "steps": [
            (1, "Run .\\scripts\\verify-thin-host.ps1 -ProjectDirectory .\\source\\ARTestCLI.", "The script prints Thin-host architecture verification: PASSED and returns 0."),
            (2, "Open source\\ARTestCLI\\ARTestCLI.vcxproj in a text editor.", "It references ARTestEngine.vcxproj and the ARTest.SDK include directory; it does not reference ARTestEngine.Core.vcxproj."),
            (3, "Search source\\ARTestCLI for ARTestEngine.Core.", "No .cpp, .h, .hpp, or .vcxproj match is found."),
            (4, "In the automated HTML report, locate StageDThinHostArchitectureTests.", "CliDependsOnlyOnThePublicSdkAndEngineDll is PASSED."),
        ],
    },
    {
        "id": "MT-D2-003",
        "title": "Offline compilation and structured diagnostics",
        "objective": "Confirm compile uses ARTestEngine.dll, never initializes hardware, and preserves exit code 3 for an invalid plan.",
        "preconditions": "Release build from MT-D2-001 completed successfully.",
        "steps": [
            (1, "Set $cli = '.\\artifacts\\bin\\x64\\Release\\ARTestCLI.exe'.", "Test-Path $cli returns True."),
            (2, "Run & $cli compile '.\\source\\Scripts\\TestScript.json'; $LASTEXITCODE.", "The CLI prints Valid script. No instruments were initialized. and returns 0; no Initialize or Shutdown text appears."),
            (3, "Run & $cli compile '.\\quality\\manual-tests\\stage-a\\data\\unknown-command.json'; $LASTEXITCODE.", "The CLI reports COMMAND_TYPE_UNKNOWN with location steps[0] and returns 3."),
            (4, "Run & $cli compile '.\\quality\\manual-tests\\stage-a\\data\\malformed-json.json'; $LASTEXITCODE.", "The CLI reports SCRIPT_JSON_INVALID, identifies the original file path, does not expose engine-api as the location, and returns 3."),
        ],
    },
    {
        "id": "MT-D2-004",
        "title": "Run command compatibility through ARTestEngine.dll",
        "objective": "Verify the legacy run contract is preserved while execution is owned by the public Engine boundary.",
        "preconditions": "$cli is defined as in MT-D2-003.",
        "steps": [
            (1, "Run & $cli run '.\\source\\Scripts\\TestScript.json'.", "The process returns exit code 0."),
            (2, "Inspect state and instrument output.", "Output shows INITIALIZING, fake CAN and power-supply initialization, RUNNING, four step records, CLEANING_UP, reverse-order shutdown, and COMPLETED."),
            (3, "Inspect the final execution summary.", "It reports PASSED with planned=4, executed=4, passed=4, failed=0, skipped=0, and attempts=4."),
        ],
    },
    {
        "id": "MT-D2-005",
        "title": "Debug control callback and continue behavior",
        "objective": "Confirm debug pauses through the API 0.2 before-step callback and the host remains responsible for user interaction.",
        "preconditions": "$cli is defined. Use an interactive PowerShell console.",
        "steps": [
            (1, "Run & $cli debug '.\\source\\Scripts\\TestScript.json'.", "The Engine initializes instruments and the CLI pauses before script step 1."),
            (2, "At Options: (n)ext, (c)ontinue, (q)uit > enter c and press Enter.", "Execution continues without another debug prompt."),
            (3, "Inspect completion and $LASTEXITCODE.", "All four steps pass, cleanup completes, terminal state is COMPLETED, and exit code is 0."),
        ],
    },
    {
        "id": "MT-D2-006",
        "title": "Zero-based breakpoint compatibility",
        "objective": "Verify break retains zero-based command-index semantics through the controlled-session callback.",
        "preconditions": "$cli is defined. Use an interactive PowerShell console.",
        "steps": [
            (1, "Run & $cli break '.\\source\\Scripts\\TestScript.json' 1.", "Script step 1 executes without a debug prompt."),
            (2, "Confirm the prompt identifies step 2: Time.WaitMs.", "The callback receives command index 1 and displays the stable script stepId 2."),
            (3, "Enter c and press Enter; then inspect $LASTEXITCODE.", "Remaining steps execute, cleanup completes, and exit code is 0."),
        ],
    },
    {
        "id": "MT-D2-007",
        "title": "Cooperative Ctrl+C cancellation through EngineClient",
        "objective": "Verify Ctrl+C reaches the public session, wakes a long command, and preserves guaranteed cleanup.",
        "preconditions": "$cli is defined. Use an interactive PowerShell console.",
        "steps": [
            (1, "Run & $cli run '.\\quality\\manual-tests\\stage-c\\data\\cancel-cleanup.json'.", "Execution reaches RUNNING and starts the long Time.WaitMs step."),
            (2, "Press Ctrl+C once while Time.WaitMs is active.", "The process does not terminate abruptly; it requests cooperative cancellation."),
            (3, "Inspect state and instrument output.", "Output shows CANCELLING, CLEANING_UP, power-supply Shutdown, and terminal CANCELLED in that order."),
            (4, "Run $LASTEXITCODE immediately after completion.", "The process exit code is 5."),
        ],
    },
    {
        "id": "MT-D2-008",
        "title": "Initialization-failure exit-code compatibility",
        "objective": "Confirm an initialization failure remains distinguishable from validation and execution failures after migration.",
        "preconditions": "$cli is defined.",
        "steps": [
            (1, "Run & $cli run '.\\quality\\manual-tests\\stage-a\\data\\initialization-failure.json'.", "The Engine enters INITIALIZING and reports POWER_SUPPLY_RESOURCE_MISSING."),
            (2, "Inspect cleanup and terminal state.", "CLEANING_UP and FAILED are emitted; no test step executes."),
            (3, "Run $LASTEXITCODE immediately after completion.", "The process exit code is 4."),
        ],
    },
    {
        "id": "MT-D2-009",
        "title": "Native extension vertical-slice regression",
        "objective": "Confirm the D1 Command Plugin to Instrument Driver flow remains functional after API 0.2 and thin-host migration.",
        "preconditions": "Release build completed and packaged extensions exist.",
        "steps": [
            (1, "Set $extensions = '.\\artifacts\\extensions\\x64\\Release'; confirm Test-Path $extensions returns True.", "The approved extension root exists."),
            (2, "Run & $cli extension-run '.\\source\\Scripts\\ExtensionScript.json' $extensions.", "The catalog loads, the sample Command Plugin invokes SimPower1, cleanup shuts down the driver, and the process returns 0."),
            (3, "Inspect the final JSON line.", "status is passed; failureKind is 0; plannedSteps, executedSteps, and passedSteps are 1."),
            (4, "In the HTML report, locate StageDEngineApiCompatibilityTests.", "NegotiatesMinorOneWithoutWritingPastItsTable is PASSED, proving API 0.1 table compatibility."),
        ],
    },
]


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_font(
        header.add_run("ARTest Quality | Stage D2 Manual Validation"),
        size=9,
        color="666666",
    )
    add_page_number(section.footer.paragraphs[0])

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_font(kicker.add_run("QUALITY ASSURANCE"), size=10, bold=True, color=BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_font(title.add_run("ARTestCLI Manual Test Report"), size=24, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(
        subtitle.add_run("Stage D2 - Thin-Host Migration"),
        size=14,
        color=DARK_BLUE,
    )

    add_metadata(
        document,
        (
            ("Document ID", "ARTEST-QA-D2-MANUAL-001"),
            ("Version", "1.0"),
            ("Status", "Pending manual execution"),
            ("Prepared for", "ARTestCLI / ARTestEngine Stage D2 acceptance"),
            ("Tester", ""),
            ("Execution date", ""),
            ("Commit / branch", ""),
        ),
    )

    document.add_heading("Purpose and acceptance scope", level=1)
    document.add_paragraph(
        "This report captures repeatable manual evidence for the Stage D2 "
        "thin-host migration. It validates the public Engine API 0.2 boundary, "
        "legacy CLI compatibility, controlled debug sessions, structured "
        "compilation, cancellation, cleanup, and D1 extension regression."
    )
    add_note(
        document,
        "Acceptance rule",
        "All nine manual cases must be Passed, and the Debug and Release "
        "automated reports must each show 61 passed, 0 failed, and 0 skipped.",
    )

    document.add_page_break()
    document.add_heading("Execution environment", level=1)
    add_metadata(
        document,
        (
            ("Operating system", "Windows:"),
            ("Visual Studio", "18 Insiders / v145"),
            ("Platform", "x64"),
            ("Configurations", "Debug and Release"),
            ("Repository", "D:\\GitHub\\main\\ARTestCLI"),
            ("PowerShell version", ""),
            ("Machine / tester", ""),
        ),
    )

    document.add_heading("Test summary", level=1)
    summary = document.add_table(rows=1, cols=4)
    for index, text in enumerate(
        ("ID", "Scenario", "Expected verdict", "Actual verdict")
    ):
        set_cell_text(summary.rows[0].cells[index], text, bold=True, color=NAVY)
    for case in TEST_CASES:
        cells = summary.add_row().cells
        set_cell_text(cells[0], case["id"], bold=True, color=NAVY)
        set_cell_text(cells[1], case["title"])
        set_cell_text(cells[2], "Passed", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], "Pending", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(summary, [1150, 5010, 1600, 1600])
    style_table(summary)

    document.add_page_break()
    for index, case in enumerate(TEST_CASES):
        document.add_heading(f'{case["id"]} | {case["title"]}', level=1)
        add_metadata(
            document,
            (
                ("Objective", case["objective"]),
                ("Preconditions", case["preconditions"]),
                ("Expected verdict", "Passed"),
            ),
        )
        document.add_heading("Procedure and expected results", level=2)
        add_steps(document, case["steps"])
        document.add_heading("Execution evidence", level=2)
        add_evidence_block(document)
        if index != len(TEST_CASES) - 1:
            document.add_page_break()

    document.add_page_break()
    document.add_heading("Final acceptance", level=1)
    add_metadata(
        document,
        (
            ("Total manual cases", str(len(TEST_CASES))),
            ("Passed", ""),
            ("Failed", ""),
            ("Blocked / Not run", ""),
            ("Automated Debug verdict", ""),
            ("Automated Release verdict", ""),
            ("Overall Stage D2 verdict", "Pending"),
        ),
    )
    document.add_heading("Approval", level=2)
    approval = document.add_table(rows=3, cols=3)
    for index, text in enumerate(("Role", "Name / Signature", "Date")):
        set_cell_text(approval.rows[0].cells[index], text, bold=True, color=NAVY)
    for row_index, role in enumerate(("Tester", "Technical reviewer"), start=1):
        set_cell_text(approval.rows[row_index].cells[0], role, bold=True)
        set_cell_text(approval.rows[row_index].cells[1], "")
        set_cell_text(approval.rows[row_index].cells[2], "")
    set_table_geometry(approval, [1900, 5060, 2400])
    style_table(approval)

    core = document.core_properties
    core.title = "ARTestCLI Manual Test Report - Stage D2"
    core.subject = "Thin-host migration acceptance"
    core.author = "ARTest Quality"
    core.keywords = "ARTestCLI, ARTestEngine, Stage D2, thin host, manual testing"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
