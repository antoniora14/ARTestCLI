from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_stage_d1_report import (
    add_evidence_block, add_metadata, add_page_number, add_steps,
    configure_styles, set_cell_text, set_font, set_table_geometry, style_table,
)

ROOT = Path(__file__).resolve().parents[2]
CASE_ROOT = ROOT / "quality/manual-tests/stage-d3.3b"
OUTPUT = CASE_ROOT / "ARTestCLI_Manual_Test_Report_Reference_Extensions_v1.0.docx"


def normalize_layout(document):
    # Keep the established ARTest form structure, with explicit monochrome
    # heading styles and visible light borders independent of Word themes.
    for name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.underline = False
        ppr = style.element.get_or_add_pPr()
        for border in list(ppr.findall(qn("w:pBdr"))):
            ppr.remove(border)
    for paragraph in document.paragraphs:
        ppr = paragraph._p.get_or_add_pPr()
        for border in list(ppr.findall(qn("w:pBdr"))):
            ppr.remove(border)
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
    for table in document.tables:
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            for key, value in (("val", "single"), ("sz", "4"), ("color", "D9D9D9")):
                element.set(qn(f"w:{key}"), value)
            borders.append(element)
        properties = table._tbl.tblPr
        old = properties.find(qn("w:tblBorders"))
        if old is not None:
            properties.remove(old)
        properties.append(borders)
        for row in table.rows:
            # Evidence may grow when a tester pastes screenshots or detailed notes.
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)


def build(output=OUTPUT):
    if output.exists():
        raise FileExistsError("Existing manual evidence is protected; choose another --output.")
    cases = json.loads((CASE_ROOT / "cases.json").read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.35)
    configure_styles(document)
    document.styles["Normal"].paragraph_format.line_spacing = 1.1
    for name in ("Heading 1", "Heading 2"):
        document.styles[name].paragraph_format.space_before = Pt(10)
        document.styles[name].paragraph_format.space_after = Pt(6)
    set_font(section.header.paragraphs[0].add_run("ARTest Quality Assurance"), size=9)
    add_page_number(section.footer.paragraphs[0])

    title = document.add_paragraph("ARTestCLI Manual Test Report", style="Title")
    title.paragraph_format.space_after = Pt(6)
    for run in title.runs:
        set_font(run, size=24, bold=True)
    document.add_paragraph("Reference Extensions and SDK Acceptance", style="Subtitle")
    document.add_paragraph(
        "Stage D3.3-B migrates the reference commands and simulated drivers to the public "
        "C++ SDK. Use this report to verify compatibility, failure handling and cooperative "
        "interruption, then record your evidence. Manual acceptance is pending."
    )
    add_metadata(document, (
        ("Document ID", "ARTEST-QA-D3.3B-MANUAL-001"),
        ("Version and status", "1.0 / Pending manual execution"),
        ("Tester and date", ""),
        ("Commit and branch", ""),
        ("Environment", "Windows / Visual Studio 18 Insiders / v145 / C++20 / x64"),
        ("SDK contracts", "SDK 0.1.1 / Engine API 0.4 / native ABI 0.1"),
    ))
    document.add_heading("Setup", 1)
    document.add_paragraph("Open PowerShell and define these variables in the same session:")
    for command in (
        r"Set-Location 'D:\GitHub\main\ARTestCLI'",
        r"$cli = '.\artifacts\bin\x64\Release\ARTestCLI.exe'",
        r"$extensions = '.\artifacts\extensions\x64\Release'",
        r"$data = '.\quality\manual-tests\stage-d3.3b\data'",
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(command)
        set_font(run, size=10)
        run.font.name = "Consolas"
    document.add_paragraph(
        "Run commands separately; inspect $LASTEXITCODE immediately after each command. "
        "extension-run prints the final JSON to the console, not to a file. debug prints "
        "a textual summary. Durations and stdout/stderr ordering can vary."
    )
    document.add_paragraph(
        "Acceptance requires all four cases to pass with recorded evidence. For negative "
        "tests, the expected nonzero exit code and diagnostics constitute a passed test. "
        "Use simulations only; do not connect physical equipment."
    )

    for case in cases:
        document.add_page_break()
        # Avoid multiline built-in headings after a page break in Word.
        identifier = document.add_paragraph(case["id"])
        identifier.paragraph_format.keep_with_next = True
        for run in identifier.runs:
            set_font(run, size=10, bold=True)
        document.add_heading(case["title"], 1)
        add_metadata(document, (
            ("Objective", case["objective"]),
            ("Preconditions", case["preconditions"]),
        ))
        document.add_heading("Procedure and expected results", 2)
        add_steps(document, [
            (number, step["action"], step["expected"])
            for number, step in enumerate(case["steps"], 1)
        ])
        set_table_geometry(document.tables[-1], [700, 4440, 4220])
        document.add_heading("Execution evidence", 2)
        add_evidence_block(document)
        document.add_paragraph(
            "Paste screenshots below or record exact log references and observed counters. "
            "Expand the evidence area as needed."
        )

    document.add_page_break()
    document.add_heading("Final acceptance", 1)
    document.add_paragraph(
        "Complete this page after executing the four cases. Do not infer manual acceptance "
        "from an automated PASSED verdict or from the expected-result column."
    )
    table = document.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Case", "Scenario", "Actual verdict")):
        set_cell_text(cell, text, bold=True)
    for case in cases:
        row = table.add_row()
        for cell, text in zip(row.cells, (case["id"], case["title"], "Pending")):
            set_cell_text(cell, text)
    set_table_geometry(table, [1900, 5560, 1900])
    style_table(table)
    document.add_paragraph()
    add_metadata(document, (
        ("Passed cases", ""),
        ("Failed cases", ""),
        ("Blocked or not run", ""),
        ("Release regression", "Expected: 161 passed / 0 failed / 0 skipped. Actual:"),
        ("Debug regression", "Expected: 161 passed / 0 failed / 0 skipped. Actual:"),
        ("Defects and evidence", ""),
        ("Overall verdict", "Pending"),
        ("Tester approval", ""),
        ("Reviewer approval", ""),
    ))
    document.add_heading("Scope of acceptance", 2)
    document.add_paragraph(
        "This validates SDK-based reference behavior, preserved identities, result-schema "
        "propagation, deadlines and cleanup. It does not certify physical drivers or freeze "
        "ABI 1.0. Automatic generation of manifests and schemas is not included."
    )
    normalize_layout(document)
    document.core_properties.title = "ARTestCLI Reference Extensions Manual Test Report"
    document.core_properties.subject = "Stage D3.3-B manual acceptance"
    document.core_properties.author = "ARTest Quality"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Create a D3.3-B report without overwriting evidence.")
    parser.add_argument("--output", type=Path, default=OUTPUT)
    build(parser.parse_args().output)
